from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título principal
title = doc.add_heading('Oracle Cloud Infrastructure - Application Validation', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtítulo
subtitle = doc.add_heading('Statement of Work', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Información General
doc.add_paragraph()
doc.add_paragraph('ISV Name: {{ isv_name }}')
doc.add_paragraph('Application Name: {{ app_name }}')
doc.add_paragraph('Version: {{ version }}')
doc.add_paragraph('Date: {{ date }}')

# Separador
doc.add_paragraph('')

# Project Participants
doc.add_heading('Project Participants', level=2)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Entity'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Role'
hdr_cells[3].text = 'Email'
# Placeholder row
row_cells = table.add_row().cells
row_cells[0].text = '{% for participant in participants %}{{ participant.entity }}'
row_cells[1].text = '{{ participant.name }}'
row_cells[2].text = '{{ participant.role }}'
row_cells[3].text = '{{ participant.email }}{% endfor %}'

doc.add_paragraph('')

# ISV Company Profile
doc.add_heading('ISV Company Profile', level=2)
doc.add_paragraph('Legal Name: {{ company_profile.legal_name }}')
doc.add_paragraph('Country of Operations: {{ company_profile.country }}')
doc.add_paragraph('Description: {{ company_profile.description }}')
doc.add_paragraph('Industry: {{ company_profile.industry }}')
doc.add_paragraph('Website: {{ company_profile.website }}')

doc.add_paragraph('')

# Project Overview
doc.add_heading('Project Overview', level=2)
doc.add_paragraph('{{ project_overview }}')

# Scope
doc.add_heading('Scope', level=2)
doc.add_paragraph('{{ scope }}')

# Major Project Milestones
doc.add_heading('Major Project Milestones', level=2)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Milestone'
hdr_cells[1].text = 'Target Date'
hdr_cells[2].text = 'Completed Date'
hdr_cells[3].text = 'Comments'
row_cells = table.add_row().cells
row_cells[0].text = '{% for milestone in milestones %}{{ milestone.name }}'
row_cells[1].text = '{{ milestone.target_date }}'
row_cells[2].text = '{{ milestone.completed_date }}'
row_cells[3].text = '{{ milestone.comments }}{% endfor %}'

# Acceptance Criteria
doc.add_heading('Acceptance Criteria', level=2)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Capability/Metric'
hdr_cells[1].text = 'Acceptance Criteria'
hdr_cells[2].text = 'Status'
hdr_cells[3].text = 'Acceptable?'
row_cells = table.add_row().cells
row_cells[0].text = '{% for criteria in acceptance_criteria %}{{ criteria.metric }}'
row_cells[1].text = '{{ criteria.description }}'
row_cells[2].text = '{{ criteria.status }}'
row_cells[3].text = '{{ criteria.acceptable }}{% endfor %}'

# Current State Architecture
doc.add_heading('Current State Architecture', level=2)
doc.add_paragraph('Diagram: {{ current_architecture.diagram_url }}')
doc.add_paragraph('Description: {{ current_architecture.description }}')

# Technology Stack
doc.add_heading('Currently Used Technology Stack', level=2)
doc.add_paragraph('{% for tech in technology_stack %}- {{ tech }}{% endfor %}')

# OCI Service Sizing
doc.add_heading('OCI Service Sizing and Amounts', level=2)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Service Name'
hdr_cells[1].text = 'Sizing Units'
hdr_cells[2].text = 'Amount'
hdr_cells[3].text = 'Comments'
row_cells = table.add_row().cells
row_cells[0].text = '{% for service in oci_services %}{{ service.name }}'
row_cells[1].text = '{{ service.units }}'
row_cells[2].text = '{{ service.amount }}'
row_cells[3].text = '{{ service.comments }}{% endfor %}'

# Future State Architecture
doc.add_heading('Future State Architecture', level=2)
doc.add_paragraph('Diagram: {{ future_architecture.diagram_url }}')
doc.add_paragraph('Deployment Overview: {{ future_architecture.overview }}')
doc.add_paragraph('Components:')
doc.add_paragraph('{% for component in future_architecture.components %}- {{ component.name }}: {{ component.description }}{% endfor %}')

# Implementation Details
doc.add_heading('Implementation Details and Configuration Settings', level=2)
doc.add_paragraph('{{ implementation_details }}')

# Closing Feedback
doc.add_heading('Closing Feedback', level=2)
doc.add_paragraph('ISV Feedback: {{ closing_feedback.isv }}')
doc.add_paragraph('Oracle Feedback: {{ closing_feedback.oracle }}')

# Footer
doc.add_paragraph('')
doc.add_paragraph('Copyright © {{ copyright_year }}, Oracle and/or its affiliates.')

# Guardar documento
doc.save('Application_Validation_Template.docx')

print("Documento generado correctamente: Application_Validation_Template.docx")

